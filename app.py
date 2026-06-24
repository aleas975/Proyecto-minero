import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from io import BytesIO
from datetime import datetime
import re

# Asegurar importaciones para exportación a Word
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("Falta instalar python-docx. Ejecuta en tu terminal: pip install python-docx")

st.set_page_config(page_title="Control de Gestión: Forecast vs Budget", layout="wide")

# =====================================================================
# SECCIÓN 1: CONFIGURACIÓN INICIAL Y CARGA DE DATOS
# =====================================================================
st.title("📊 Control de Gestión Minera: Forecast vs Budget")
st.markdown("Plataforma corporativa para el análisis y visualización de desviaciones presupuestarias.")

with st.sidebar:
    st.header("📂 Ingesta de Datos")
    uploaded_file = st.file_uploader("Sube el archivo Excel del Proyecto", type=["xlsx"])

if uploaded_file:
    try:
        xls = pd.ExcelFile(uploaded_file)
        hojas_disponibles = xls.sheet_names

        st.sidebar.markdown("### 🗺️ Mapeo de Estructura")
        hoja_b = st.sidebar.selectbox("Selecciona la hoja de BUDGET", hojas_disponibles, index=0)
        hoja_f = st.sidebar.selectbox("Selecciona la hoja de FORECAST", hojas_disponibles,
                                      index=min(1, len(hojas_disponibles) - 1))

        df_forecast = pd.read_excel(uploaded_file, sheet_name=hoja_f)
        df_budget = pd.read_excel(uploaded_file, sheet_name=hoja_b)


        # 🛡️ INTERCEPTOR ULTRA-ROBUSTO DE FECHAS Y FORMATOS DE EXCEL
        def estandarizar_columnas(df):
            dict_meses_es_en = {
                'jan': 'Jan', 'feb': 'Feb', 'mar': 'Mar', 'apr': 'Apr', 'may': 'May', 'jun': 'Jun',
                'jul': 'Jul', 'aug': 'Aug', 'sep': 'Sep', 'oct': 'Oct', 'nov': 'Nov', 'dec': 'Dec',
                'ene': 'Jan', 'abr': 'Apr', 'ago': 'Aug', 'dic': 'Dec'
            }
            abrev_en = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
            nuevas_cols = []

            for c in df.columns:
                # Caso 1: Es un objeto fecha nativo de Pandas/Python
                if isinstance(c, (pd.Timestamp, datetime)):
                    mes_en = abrev_en[c.month - 1]
                    nuevas_cols.append(f"{mes_en}-{str(c.year)[-2:]}")
                    continue

                # Limpieza inicial de texto y espacios fantasmas
                str_c = str(c).strip().replace('\xa0', '')

                # Caso 2: Es un string con formato fecha oculto (ej: "2026-01-01")
                try:
                    dt = pd.to_datetime(str_c, errors='coerce')
                    if not pd.isna(dt) and (2000 < dt.year < 2100):
                        mes_en = abrev_en[dt.month - 1]
                        nuevas_cols.append(f"{mes_en}-{str(dt.year)[-2:]}")
                        continue
                except:
                    pass

                # Caso 3: Es un string tipo "Jan-26" o "Ene-26" con o sin espacios
                encontrado = False
                str_c_lower = str_c.lower()
                for k, v in dict_meses_es_en.items():
                    if k in str_c_lower:
                        num_anios = re.findall(r'\d+', str_c)
                        if num_anios:
                            anio_2d = num_anios[-1][-2:]  # Toma los últimos 2 dígitos del año
                            nuevas_cols.append(f"{v}-{anio_2d}")
                            encontrado = True
                            break

                if not encontrado:
                    nuevas_cols.append(str_c)

            df.columns = nuevas_cols
            return df


        # Aplicamos el escudo extractor a ambos DataFrames
        df_forecast = estandarizar_columnas(df_forecast)
        df_budget = estandarizar_columnas(df_budget)

        # Estandarización de nomenclaturas de portales (Gerencia -> Trabajador)
        for df in [df_forecast, df_budget]:
            if 'Gerencia' in df.columns:
                df['Gerencia'] = df['Gerencia'].astype(str).str.strip().replace(
                    {'Operaciones': 'Trabajador', 'OPERACIONES': 'TRABAJADOR'})

        # =====================================================================
        # SECCIÓN 2: PROCESAMIENTO ANALÍTICO Y UNIÓN DE BASES
        # =====================================================================
        abrev_meses = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']

        # Búsqueda segura usando el prefijo estandarizado exacto "Mes-" (ej: "Jan-")
        meses_f_map = {m: [c for c in df_forecast.columns if c.startswith(f"{m}-")][0] for m in abrev_meses if
                       [c for c in df_forecast.columns if c.startswith(f"{m}-")]}
        meses_b_map = {m: [c for c in df_budget.columns if c.startswith(f"{m}-")][0] for m in abrev_meses if
                       [c for c in df_budget.columns if c.startswith(f"{m}-")]}
        meses_validos = [m for m in abrev_meses if m in meses_f_map and m in meses_b_map]

        if not meses_validos:
            st.error(
                "🚨 Error Crítico: No se detectaron las columnas de los meses en tu Excel. Asegúrate de que las hojas contengan los encabezados de los meses (ej: Jan-26, Feb-26).")
        else:
            columnas_meses_f = [meses_f_map[m] for m in meses_validos]
            columnas_meses_b = [meses_b_map[m] for m in meses_validos]

            anio_detectado = "".join([char for char in columnas_meses_b[0] if char.isdigit()])
            if len(anio_detectado) > 2:
                anio_detectado = anio_detectado[-2:]

            col_fy_budget = f"FY{anio_detectado}"
            if col_fy_budget not in df_budget.columns:
                col_fy_budget = f"FY20{anio_detectado}"
                if col_fy_budget not in df_budget.columns:
                    cols_fy_fallback = [c for c in df_budget.columns if c.startswith('FY')]
                    col_fy_budget = cols_fy_fallback[0] if cols_fy_fallback else None

            cols_fy_todas = sorted([c for c in df_budget.columns if c.startswith('FY')])

            if 'CC' in df_forecast.columns and 'CC' in df_budget.columns:
                llaves_cruce = ['CC']
                if 'Classif' in df_forecast.columns and 'Classif' in df_budget.columns:
                    llaves_cruce.append('Classif')
                if 'Gerencia' in df_forecast.columns and 'Gerencia' in df_budget.columns:
                    llaves_cruce.append('Gerencia')
                if 'Desc Item' in df_forecast.columns and 'Desc Item' in df_budget.columns:
                    llaves_cruce.append('Desc Item')

                cols_b_extra = list(set(columnas_meses_b.copy() + cols_fy_todas))

                df_f_melt = df_forecast.groupby(llaves_cruce)[columnas_meses_f].sum().reset_index()
                df_b_melt = df_budget.groupby(llaves_cruce)[cols_b_extra].sum().reset_index()

                meses_b_rename = {meses_b_map[m]: f"{m}_Bud" for m in meses_validos}
                df_b_melt = df_b_melt.rename(columns=meses_b_rename)

                df_merged = pd.merge(df_f_melt, df_b_melt, on=llaves_cruce, how='outer')

                columnas_numericas = columnas_meses_f + [f"{m}_Bud" for m in meses_validos] + cols_fy_todas
                df_merged[columnas_numericas] = df_merged[columnas_numericas].fillna(0)

                # =====================================================================
                # SECCIÓN 3: MOTOR PREDICTIVO DINÁMICO NO LINEAL (X+Y)
                # =====================================================================
                st.sidebar.markdown("### ⚙️ Filtros Operacionales")
                filtro_classif = st.sidebar.selectbox("Clasificación de Costo (Classif)",
                                                      ["Todas"] + list(df_merged['Classif'].dropna().unique()))
                if filtro_classif != "Todas":
                    df_merged = df_merged[df_merged['Classif'] == filtro_classif]

                filtro_gerencia = st.sidebar.selectbox("Gerencia",
                                                       ["Todas"] + list(df_merged['Gerencia'].dropna().unique()))
                if filtro_gerencia != "Todas":
                    df_merged = df_merged[df_merged['Gerencia'] == filtro_gerencia]

                st.sidebar.markdown("### 🧠 Motor Predictivo Dinámico (X+Y)")
                meses_reales_n = st.sidebar.slider("Meses Reales Consolidados (Corte temporal)",
                                                   min_value=1, max_value=len(meses_validos) - 1, value=5)

                cuentas_fijas = ['Labor', 'Maintenance', 'Spare Parts', 'Expenses']

                # Fase 1: Ventana real con Fallback inteligente a Budget ante ceros
                for i in range(meses_reales_n):
                    m = meses_validos[i]
                    col_f = meses_f_map[m]
                    col_b = f"{m}_Bud"
                    col_calc = f"{m}_Calc"
                    df_merged[col_calc] = np.where((df_merged[col_f] == 0) | (df_merged[col_f].isna()),
                                                   df_merged[col_b],
                                                   df_merged[col_f])

                # Fase 2: Factor de ejecución YTD por fila independiente (Evita compensación cruzada)
                cols_reales_ytd = [f"{meses_validos[j]}_Calc" for j in range(meses_reales_n)]
                cols_bud_ytd = [f"{meses_validos[j]}_Bud" for j in range(meses_reales_n)]
                sum_real_ytd = df_merged[cols_reales_ytd].sum(axis=1)
                sum_bud_ytd = df_merged[cols_bud_ytd].sum(axis=1)
                factor_ejecucion = np.where(sum_bud_ytd > 0, sum_real_ytd / sum_bud_ytd, 1.0)
                factor_ejecucion = np.clip(factor_ejecucion, 0.85, 1.15)

                # Fase 3: Proyección del Futuro (No lineal, respeta estacionalidad del presupuesto)
                for i in range(meses_reales_n, len(meses_validos)):
                    m = meses_validos[i]
                    col_b = f"{m}_Bud"
                    col_calc = f"{m}_Calc"
                    condicion_fija = df_merged['Classif'].isin(cuentas_fijas)
                    df_merged[col_calc] = np.where(condicion_fija, df_merged[col_b],
                                                   df_merged[col_b] * factor_ejecucion)

                # Calcular Varianza Total por Fila para Reportes
                df_merged['Varianza_Total_Fila'] = 0.0
                for m in meses_validos:
                    df_merged['Varianza_Total_Fila'] += (df_merged[f"{m}_Calc"] - df_merged[f"{m}_Bud"])

                # =====================================================================
                # SECCIÓN 4: CÁLCULOS GLOBALES
                # =====================================================================
                valores_forecast_full = [df_merged[f"{m}_Calc"].sum() for m in meses_validos]
                valores_budget_full_list = [df_merged[f"{m}_Bud"].sum() for m in meses_validos]

                if col_fy_budget and col_fy_budget in df_merged.columns:
                    sum_budget_m = df_merged[col_fy_budget].sum() / 1_000_000
                else:
                    sum_budget_m = sum(valores_budget_full_list) / 1_000_000

                sum_forecast_m = sum(valores_forecast_full) / 1_000_000
                varianza_total_m = sum_forecast_m - sum_budget_m
                varianzas_mensuales_full = [f - b for f, b in zip(valores_forecast_full, valores_budget_full_list)]

                tab_anual, tab_temporal, tab_multi_presupuesto, tab_export = st.tabs([
                    f"📅 Análisis Consolidado (Forecast {meses_reales_n}+{len(meses_validos) - meses_reales_n})",
                    "⏳ Semáforo de Control YTD",
                    "📈 Tendencias Plurianuales",
                    "📥 Exportar Reportes"
                ])

                # =====================================================================
                # SECCIÓN 5: INTERFAZ - PESTAÑA ANUAL (FULL YEAR)
                # =====================================================================
                with tab_anual:
                    st.markdown(f"### 🎯 Desempeño Financiero Anual Consolidado")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Forecast Anual Proyectado (FY)", f"$ {sum_forecast_m:,.2f} M")
                    col2.metric(f"Budget Anual Oficial Corregido", f"$ {sum_budget_m:,.2f} M")
                    col3.metric("Varianza Total Anual", f"$ {varianza_total_m:,.2f} M",
                                delta=f"$ {varianza_total_m:,.2f} M",
                                delta_color="inverse" if varianza_total_m >= 0 else "normal")
                    st.divider()
                    fig_comb = go.Figure()
                    fig_comb.add_trace(
                        go.Bar(x=meses_validos, y=valores_budget_full_list, name="Budget", marker_color="#1f77b4",
                               opacity=0.65))
                    fig_comb.add_trace(
                        go.Bar(x=meses_validos, y=valores_forecast_full, name=f"Forecast", marker_color="#ff7f0e",
                               opacity=0.85))
                    fig_comb.add_trace(
                        go.Scatter(x=meses_validos, y=varianzas_mensuales_full, name="Varianza", mode="lines+markers",
                                   line=dict(color="#d62728", width=3), yaxis="y2"))
                    fig_comb.update_layout(barmode="group", hovermode="x unified",
                                           title="Distribución de Volúmenes y Variación del Periodo",
                                           legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                                           yaxis=dict(title="Montos Absolutos ($)", showgrid=True),
                                           yaxis2=dict(title="Varianza Neta ($)", overlaying="y", side="right",
                                                       showgrid=False), height=500)
                    fig_comb.add_vline(x=meses_reales_n - 0.5, line_width=2, line_dash="dash", line_color="gray",
                                       annotation_text="Corte Proyección")
                    st.plotly_chart(fig_comb, use_container_width=True)

                # =====================================================================
                # SECCIÓN 6: INTERFAZ - SEMÁFORO DE GESTIÓN (YTD)
                # =====================================================================
                with tab_temporal:
                    mes_corte = st.selectbox("Seleccionar mes de corte (Gasto acumulado hasta hoy):", meses_validos,
                                             index=meses_reales_n - 1)
                    idx_corte = meses_validos.index(mes_corte)
                    meses_hasta_hoy = meses_validos[:idx_corte + 1]
                    valores_forecast_hoy = [df_merged[f"{m}_Calc"].sum() for m in meses_hasta_hoy]
                    valores_budget_hoy = [df_merged[f"{m}_Bud"].sum() for m in meses_hasta_hoy]
                    varianzas_mensuales_hoy = [f - b for f, b in zip(valores_forecast_hoy, valores_budget_hoy)]
                    sum_forecast_hoy_m = sum(valores_forecast_hoy) / 1_000_000
                    sum_budget_hoy_m = sum(valores_budget_hoy) / 1_000_000
                    varianza_hoy_m = sum_forecast_hoy_m - sum_budget_hoy_m
                    pct_varianza_ytd = (varianza_hoy_m / sum_budget_hoy_m * 100) if sum_budget_hoy_m > 0 else 0

                    col_hoy1, col_hoy2, col_hoy3 = st.columns(3)
                    col_hoy1.metric(f"Forecast Acumulado (Ene-{mes_corte})", f"$ {sum_forecast_hoy_m:,.2f} M")
                    col_hoy2.metric(f"Budget Acumulado (Ene-{mes_corte})", f"$ {sum_budget_hoy_m:,.2f} M")
                    col_hoy3.metric(f"Varianza Acumulada (Hasta {mes_corte})", f"$ {varianza_hoy_m:,.2f} M",
                                    delta=f"$ {varianza_hoy_m:,.2f} M",
                                    delta_color="inverse" if varianza_hoy_m >= 0 else "normal")

                    st.divider()
                    st.markdown("#### 🚦 Semáforo de Control Presupuestario (YTD)")
                    if pct_varianza_ytd <= 0:
                        st.success(
                            f"🟢 **ÓPTIMO:** Ahorro o cumplimiento perfecto ({pct_varianza_ytd:+.1f}%). El gasto acumulado es menor o igual al presupuesto autorizado.")
                    elif 0 < pct_varianza_ytd <= 5:
                        st.warning(
                            f"🟡 **PRECAUCIÓN:** Desviación leve ({pct_varianza_ytd:+.1f}%). El gasto se encuentra dentro del margen de tolerancia operativo.")
                    else:
                        st.error(
                            f"🔴 **ALERTA CRÍTICA:** Sobregasto detectado ({pct_varianza_ytd:+.1f}%). Requiere plan de acción inmediato.")
                    st.divider()
                    fig_water_hoy = go.Figure(
                        go.Waterfall(name="Variación Acumulada", orientation="v", x=meses_hasta_hoy,
                                     textposition="outside",
                                     text=[f"${v / 1_000_000:.2f}M" for v in varianzas_mensuales_hoy],
                                     y=varianzas_mensuales_hoy,
                                     connector=dict(line=dict(color="rgb(63, 63, 63)", width=1.5)),
                                     decreasing=dict(marker=dict(color="#2ca02c")),
                                     increasing=dict(marker=dict(color="#d62728"))))
                    fig_water_hoy.update_layout(title=f"Gráfico de Cascada Acumulado ({mes_corte})", showlegend=False,
                                                height=550)
                    st.plotly_chart(fig_water_hoy, use_container_width=True)

                # =====================================================================
                # SECCIÓN 7: INTERFAZ - MULTI-PRESUPUESTO HORIZONTE PLURIANUAL
                # =====================================================================
                with tab_multi_presupuesto:
                    st.markdown("### 📈 Análisis de Evolución Presupuestaria Horizonte Plurianual")
                    valores_multi_fy = [df_merged[c].sum() / 1_000_000 for c in cols_fy_todas]
                    fig_trend_fy = go.Figure(go.Scatter(x=cols_fy_todas, y=valores_multi_fy, mode="lines+markers+text",
                                                        text=[f"${v:,.2f}M" for v in valores_multi_fy],
                                                        textposition="top center", line=dict(color="#9467bd", width=4),
                                                        marker=dict(size=10, symbol="diamond")))
                    fig_trend_fy.update_layout(title="Tendencia de Presupuestos Consolidados por Año de Ejercicio",
                                               xaxis_title="Periodos", yaxis_title="Monto (M$)", height=400)
                    st.plotly_chart(fig_trend_fy, use_container_width=True)

                # =====================================================================
                # SECCIÓN 8: EXPORTAR REPORTES AUTOMATIZADOS (WORD Y EXCEL)
                # =====================================================================
                with tab_export:
                    st.markdown("### 📥 Generación de Reportes Ejecutivos")
                    st.markdown(
                        "Descarga el detalle matemático en Excel o el informe consolidado con hallazgos estratégicos en Word.")

                    # 1. GENERADOR EXCEL SÁBANA ANALÍTICA
                    output_excel = BytesIO()
                    with pd.ExcelWriter(output_excel, engine="xlsxwriter") as writer:
                        columnas_export = [c for c in ['CC', 'Classif', 'Gerencia', 'Desc Item'] if
                                           c in df_merged.columns]
                        columnas_export += [f"{m}_Calc" for m in meses_validos] + [f"{m}_Bud" for m in
                                                                                   meses_validos] + [
                                               'Varianza_Total_Fila']
                        df_excel = df_merged[columnas_export].sort_values(by='Varianza_Total_Fila', ascending=False)
                        df_excel.to_excel(writer, sheet_name="Forecast_Detalle", index=False)


                    # 2. GENERADOR WORD CON REDACCIÓN DINÁMICA DE HALLAZGOS
                    def generar_word():
                        doc = Document()
                        titulo = doc.add_heading(
                            f'REPORTE EJECUTIVO DE FORECAST OPERATIVO ({meses_reales_n}+{len(meses_validos) - meses_reales_n})',
                            0)
                        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph(
                            f"Generado el: {datetime.now().strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

                        doc.add_heading('1. Resumen Financiero Global', level=1)
                        doc.add_paragraph(f"El presupuesto oficial proyectado es de $ {sum_budget_m:,.2f} Millones.")
                        doc.add_paragraph(
                            f"El modelo predictivo Forecast {meses_reales_n}+{len(meses_validos) - meses_reales_n} estima un cierre anual de $ {sum_forecast_m:,.2f} Millones.")

                        estado = "Ahorro/Sub-ejecución" if varianza_total_m <= 0 else "Sobregasto"
                        doc.add_paragraph(
                            f"Esto representa una desviación neta de $ {varianza_total_m:,.2f} Millones ({estado}).")

                        doc.add_heading('2. Estado de Alerta (Semáforo YTD)', level=1)
                        if pct_varianza_ytd <= 0:
                            doc.add_paragraph(
                                f"ESTADO VERDE (Óptimo): El gasto acumulado presenta una eficiencia del {abs(pct_varianza_ytd):.2f}%.")
                        elif 0 < pct_varianza_ytd <= 5:
                            doc.add_paragraph(
                                f"ESTADO AMARILLO (Precaución): Desviación del {pct_varianza_ytd:.2f}%, dentro de las bandas de flotación aceptadas.")
                        else:
                            doc.add_paragraph(
                                f"ESTADO ROJO (Alerta): Sobregasto crítico del {pct_varianza_ytd:.2f}%. Se exige mitigación inmediata.")

                        doc.add_heading('3. Hallazgos Estratégicos y Lógicas Aplicadas', level=1)
                        doc.add_paragraph("Regla 1 - Protección de Cuentas Inelásticas:", style='List Bullet')
                        doc.add_paragraph(
                            "Las cuentas fijas (Labor, Maintenance, Spare Parts) fueron bloqueadas a nivel de presupuesto oficial para los meses proyectados, previniendo falsas volatilidades en sueldos y contratos a suma alzada.")
                        doc.add_paragraph("Regla 2 - Proyección Variable (Run-Rate Controlado):", style='List Bullet')
                        doc.add_paragraph(
                            "Las cuentas como Fuel y Power fueron proyectadas respetando su tendencia de consumo real (YTD), pero aplicando un techo analítico de +/- 15% para evitar distorsiones por eventos únicos.")
                        doc.add_paragraph("Regla 3 - Cero Compensación Cruzada:", style='List Bullet')
                        doc.add_paragraph(
                            "El factor de desviación se calculó independientemente por cada centro de costo y clasificación, asegurando que un ahorro en contratistas no financie un sobregasto en energía de forma invisible.")

                        df_agrupado = df_merged.groupby('Classif')['Varianza_Total_Fila'].sum().sort_values()
                        mayor_ahorro = df_agrupado.index[0] if df_agrupado.iloc[0] < 0 else "Ninguno"
                        mayor_gasto = df_agrupado.index[-1] if df_agrupado.iloc[-1] > 0 else "Ninguno"

                        doc.add_heading('4. Focos de Atención Directa', level=1)
                        doc.add_paragraph(
                            f"Mayor foco de sobregasto actual: La cuenta '{mayor_gasto}' lidera las varianzas negativas del periodo.")
                        doc.add_paragraph(
                            f"Principal generador de ahorro: La cuenta '{mayor_ahorro}' está absorbiendo los impactos operacionales permitiendo eficiencias netas.")

                        output_word = BytesIO()
                        doc.save(output_word)
                        output_word.seek(0)
                        return output_word


                    col_exp1, col_exp2 = st.columns(2)
                    with col_exp1:
                        st.download_button(
                            label="📊 Descargar Sábana Analítica (Excel)",
                            data=output_excel.getvalue(),
                            file_name=f"Forecast_{meses_reales_n}plus{len(meses_validos) - meses_reales_n}_Detalle.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    with col_exp2:
                        try:
                            word_file = generar_word()
                            st.download_button(
                                label="📄 Descargar Informe y Hallazgos (Word)",
                                data=word_file,
                                file_name=f"Reporte_Forecast_{meses_reales_n}plus{len(meses_validos) - meses_reales_n}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        except Exception as e:
                            st.error(f"Error generando Word: {e}")

                # MATRIZ INFERIOR DETALLADA
                st.subheader("📋 Matriz Detallada por Centro de Costo")
                columnas_visibles = [c for c in ['CC', 'Classif', 'Gerencia', 'Desc Item'] if
                                     c in df_merged.columns] + ['Varianza_Total_Fila']
                st.dataframe(df_merged[columnas_visibles].sort_values(by='Varianza_Total_Fila', ascending=False),
                             use_container_width=True)

            else:
                st.error("Error estructural: No se encontró la columna clave 'CC' en ambas hojas.")

    except Exception as e:
        st.error(f"Error crítico de procesamiento: {str(e)}")
else:
    st.info(
        "👈 Esperando carga de la planilla de control de gestión minera (.xlsx) para inicializar el pipeline analítico.")